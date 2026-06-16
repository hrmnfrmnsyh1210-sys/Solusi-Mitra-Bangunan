# -*- coding: utf-8 -*-
"""Generate dokumen Word Bab IV (4.4 - 4.7) Sistem Manajemen Inventori."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- Default style ----------
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:color'), 'BFBFBF')
        tcBorders.append(el)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def h1(text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x31, 0x2E, 0x81)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def para(text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(2)
    return p


def code_block(caption, code):
    """Potongan kode di dalam tabel 1 sel dengan latar abu-abu, font monospace."""
    if caption:
        cap = doc.add_paragraph()
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        cap.paragraph_format.space_after = Pt(2)
        cap.paragraph_format.space_before = Pt(6)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, 'F5F5F5')
    set_cell_border(cell)

    cell.paragraphs[0].text = ''
    first = True
    for line in code.rstrip('\n').split('\n'):
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line != '' else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def screenshot(caption):
    """Kotak placeholder untuk menempelkan screenshot."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, 'FAFAFA')
    set_cell_border(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run('[  Tempelkan screenshot di sini  ]')
    r.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r.font.size = Pt(11)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(caption)
    rc.italic = True
    rc.font.size = Pt(10.5)
    rc.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    cap.paragraph_format.space_after = Pt(10)


def data_table(headers, rows):
    """Tabel data biasa (bukan kode) dengan header berwarna."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        shade_cell(hdr[i], '4F46E5')
        set_cell_border(hdr[i])
        p = hdr[i].paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_border(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


# ============================================================
# 4.3 MEMBUAT DATABASE
# ============================================================
h1('4.3  Membuat Database')

para('Database merupakan tempat penyimpanan seluruh data aplikasi. Sistem Manajemen '
     'Inventori ini menggunakan DBMS MySQL dengan nama database manajemen_inventori. '
     'Pembuatan tabel tidak dilakukan secara manual melalui phpMyAdmin, melainkan '
     'menggunakan fitur Migration pada CodeIgniter 4. Dengan migration, struktur tabel '
     'ditulis dalam bentuk kode PHP sehingga mudah dibuat ulang, terdokumentasi, dan '
     'konsisten di setiap komputer.')

para('Konfigurasi koneksi database diatur pada berkas .env. Selanjutnya seluruh tabel '
     'dibuat dengan satu perintah migrasi, dan data awal (kategori, supplier, barang '
     'contoh, serta akun pengguna) diisi melalui Seeder.')

h2('a. Struktur Tabel')
para('Database manajemen_inventori terdiri dari beberapa tabel utama yang saling '
     'berelasi, antara lain:', justify=False)

data_table(
    ['No', 'Nama Tabel', 'Fungsi'],
    [
        ['1', 'users', 'Menyimpan data akun pengguna (login admin).'],
        ['2', 'kategori', 'Menyimpan kategori barang.'],
        ['3', 'supplier', 'Menyimpan data pemasok barang.'],
        ['4', 'barang', 'Data master barang (kode, nama, harga, stok).'],
        ['5', 'stok_masuk', 'Header transaksi barang masuk dari supplier.'],
        ['6', 'detail_stok_masuk', 'Rincian item pada tiap transaksi barang masuk.'],
        ['7', 'penjualan', 'Header transaksi penjualan.'],
        ['8', 'detail_penjualan', 'Rincian item pada tiap transaksi penjualan.'],
        ['9', 'retur_penjualan', 'Data retur / pengembalian barang dari penjualan.'],
        ['10', 'penyesuaian_stok', 'Penyesuaian stok (koreksi stok barang).'],
        ['11', 'customers', 'Data pelanggan (untuk fitur toko/shop).'],
    ]
)

h2('b. Potongan Kode')

code_block('Konfigurasi koneksi database — berkas .env',
'''# DATABASE
database.default.hostname = localhost
database.default.database = manajemen_inventori
database.default.username = root
database.default.password =
database.default.DBDriver = MySQLi
database.default.port     = 3306''')

code_block('Contoh Migration tabel barang — app/Database/Migrations/..._CreateBarangTable.php',
'''public function up()
{
    $this->forge->addField([
        'id' => [
            'type' => 'INT', 'constraint' => 11,
            'unsigned' => true, 'auto_increment' => true,
        ],
        'kode_barang'  => ['type' => 'VARCHAR', 'constraint' => 50, 'unique' => true],
        'nama_barang'  => ['type' => 'VARCHAR', 'constraint' => 150],
        'id_kategori'  => ['type' => 'INT', 'constraint' => 11, 'unsigned' => true, 'null' => true],
        'satuan'       => ['type' => 'VARCHAR', 'constraint' => 30],
        'harga_beli'   => ['type' => 'DECIMAL', 'constraint' => '15,2', 'default' => 0],
        'harga_jual'   => ['type' => 'DECIMAL', 'constraint' => '15,2', 'default' => 0],
        'stok'         => ['type' => 'INT', 'constraint' => 11, 'default' => 0],
        'stok_minimum' => ['type' => 'INT', 'constraint' => 11, 'default' => 0],
        'keterangan'   => ['type' => 'TEXT', 'null' => true],
        'created_at'   => ['type' => 'DATETIME', 'null' => true],
        'updated_at'   => ['type' => 'DATETIME', 'null' => true],
    ]);
    $this->forge->addPrimaryKey('id');
    // Relasi ke tabel kategori
    $this->forge->addForeignKey('id_kategori', 'kategori', 'id', 'SET NULL', 'CASCADE');
    $this->forge->createTable('barang');
}''')

code_block('Contoh relasi antar tabel (detail_penjualan)',
'''$this->forge->addPrimaryKey('id');
$this->forge->addForeignKey('id_penjualan', 'penjualan', 'id', 'CASCADE', 'CASCADE');
$this->forge->addForeignKey('id_barang',    'barang',    'id', 'CASCADE', 'CASCADE');
$this->forge->createTable('detail_penjualan');''')

code_block('Perintah membuat tabel & mengisi data awal (terminal)',
'''# Jalankan seluruh migration untuk membuat semua tabel
php spark migrate

# Isi data awal (kategori, supplier, barang, akun login)
php spark db:seed DatabaseSeeder''')

h2('c. Penjelasan')
bullet('Database manajemen_inventori dibuat pada MySQL, dengan koneksi dikonfigurasi di berkas .env.')
bullet('Setiap tabel didefinisikan dalam sebuah file Migration menggunakan $this->forge, lalu dibuat otomatis dengan perintah php spark migrate.')
bullet('Relasi antar tabel dibentuk dengan addForeignKey(), misalnya detail_penjualan terhubung ke penjualan dan barang. Aturan ON DELETE CASCADE memastikan data rinci ikut terhapus bila induknya dihapus.')
bullet('Kolom harga menggunakan tipe DECIMAL(15,2) agar nilai uang akurat, sedangkan kode_barang dan no_transaksi bersifat unique untuk mencegah duplikasi.')
bullet('Data awal diisi melalui Seeder (php spark db:seed) sehingga aplikasi langsung memiliki akun login dan data contoh.')

h2('d. Screenshot')
screenshot('Gambar 4.3.1  Struktur Database manajemen_inventori di phpMyAdmin')
screenshot('Gambar 4.3.2  Hasil Perintah Migrasi (php spark migrate)')

# ============================================================
# 4.4 MEMBUAT LOGIN
# ============================================================
h1('4.4  Membuat Login')

para('Fitur login berfungsi sebagai pintu masuk (autentikasi) sistem agar hanya '
     'pengguna terdaftar yang dapat mengakses dasbor pengelolaan inventori. Proses '
     'autentikasi dibangun mengikuti pola MVC pada framework CodeIgniter 4, yaitu '
     'Route untuk mengarahkan permintaan, Controller (AuthController) untuk memproses '
     'login, Model (UserModel) untuk mengambil data pengguna dari basis data, dan View '
     '(auth/login) sebagai tampilan formulir.')

para('Mekanisme keamanan yang diterapkan: (1) password disimpan dalam bentuk hash '
     'sehingga tidak pernah disimpan sebagai teks asli, lalu diverifikasi dengan fungsi '
     'password_verify(); (2) status login disimpan dalam session; dan (3) setiap form '
     'dilindungi token CSRF melalui csrf_field().')

h2('a. Potongan Kode')

code_block('Route — app/Config/Routes.php',
'''// Auth
$routes->get('login',  'AuthController::login');
$routes->post('login', 'AuthController::doLogin');
$routes->get('logout', 'AuthController::logout');''')

code_block('Controller — app/Controllers/AuthController.php',
'''public function doLogin()
{
    $username = $this->request->getPost('username');
    $password = $this->request->getPost('password');

    $model = new UserModel();
    $user  = $model->findByUsername($username);

    // Verifikasi password yang sudah di-hash
    if ($user && password_verify($password, $user['password'])) {
        session()->set([
            'logged_in' => true,
            'user_id'   => $user['id'],
            'nama'      => $user['nama'],
            'username'  => $user['username'],
        ]);
        return redirect()->to(base_url('dashboard'));
    }

    return redirect()->back()
        ->with('error', 'Username atau password salah.')
        ->withInput();
}''')

code_block('Model — app/Models/UserModel.php',
'''public function findByUsername(string $username): array|null
{
    return $this->where('username', $username)->first();
}''')

code_block('View — app/Views/auth/login.php (potongan form)',
'''<form action="<?= base_url('login') ?>" method="post">
    <?= csrf_field() ?>
    <input type="text" name="username" placeholder="Masukkan username Anda" required>
    <input type="password" name="password" placeholder="Masukkan password Anda" required>
    <button type="submit">Masuk ke Sistem</button>
</form>''')

h2('b. Penjelasan')
bullet('Method login() menampilkan halaman formulir, dan langsung mengalihkan ke dashboard jika pengguna sudah login.')
bullet('Method doLogin() mengambil username & password dari form, mencari pengguna melalui UserModel::findByUsername(), lalu memverifikasi password dengan password_verify().')
bullet('Jika cocok, data pengguna disimpan ke session (logged_in = true) dan diarahkan ke dashboard. Jika gagal, pengguna dikembalikan dengan pesan error.')
bullet('Method logout() menghapus seluruh session dengan session()->destroy() dan mengarahkan kembali ke halaman login.')

h2('c. Screenshot')
screenshot('Gambar 4.4  Tampilan Halaman Login')

# ============================================================
# 4.5 CRUD BARANG
# ============================================================
h1('4.5  Membuat CRUD Barang')

para('CRUD (Create, Read, Update, Delete) Barang merupakan fitur inti untuk mengelola '
     'data master barang. Melalui fitur ini pengguna dapat menambah, menampilkan, '
     'mengubah, dan menghapus data barang, termasuk mengunggah gambar produk. Data '
     'barang dikelola oleh BarangController dan BarangModel, dengan validasi serta '
     'pembuatan kode barang otomatis (BRG-001, BRG-002, dst.).')

h2('a. Potongan Kode')

code_block('Route — app/Config/Routes.php',
'''$routes->get('barang',               'BarangController::index');
$routes->get('barang/create',        'BarangController::create');
$routes->post('barang/store',        'BarangController::store');
$routes->get('barang/edit/(:num)',   'BarangController::edit/$1');
$routes->post('barang/update/(:num)', 'BarangController::update/$1');
$routes->get('barang/delete/(:num)', 'BarangController::delete/$1');''')

code_block('Model — aturan validasi & kode otomatis (BarangModel.php)',
'''protected $validationRules = [
    'kode_barang' => 'required|max_length[50]|is_unique[barang.kode_barang,id,{id}]',
    'nama_barang' => 'required|min_length[2]|max_length[150]',
    'satuan'      => 'required|max_length[30]',
    'harga_beli'  => 'required|decimal',
    'harga_jual'  => 'required|decimal',
    'stok'        => 'required|integer',
];

// Generate kode barang otomatis: BRG-001, BRG-002, ...
public function generateKode()
{
    $last = $this->select('kode_barang')->orderBy('id', 'DESC')->first();
    if (!$last) return 'BRG-001';
    $num = (int) substr($last['kode_barang'], 4);
    return 'BRG-' . str_pad($num + 1, 3, '0', STR_PAD_LEFT);
}''')

code_block('Controller — READ (index) dengan pencarian, filter & paginasi',
'''public function index()
{
    $search   = $this->request->getGet('search');
    $kategori = $this->request->getGet('kategori');

    $builder = $this->model->select('barang.*, kategori.nama_kategori')
                           ->join('kategori', 'kategori.id = barang.id_kategori', 'left')
                           ->orderBy('barang.nama_barang');

    if ($search) {
        $builder->groupStart()
                ->like('barang.nama_barang', $search)
                ->orLike('barang.kode_barang', $search)
                ->groupEnd();
    }
    if ($kategori) $builder->where('barang.id_kategori', $kategori);

    $barangs = $builder->paginate(10, 'default');
    return view('barang/index', ['barangs' => $barangs, ...]);
}''')

code_block('Controller — CREATE (store) dengan upload gambar',
'''public function store()
{
    $data   = $this->request->getPost();
    $gambar = $this->handleUploadGambar();   // validasi tipe & ukuran file
    if ($gambar !== null) $data['gambar'] = $gambar;

    if (!$this->model->save($data)) {
        return redirect()->back()->withInput()
            ->with('error', implode('<br>', $this->model->errors()));
    }
    return redirect()->to('barang')->with('success', 'Barang berhasil ditambahkan.');
}''')

code_block('Controller — UPDATE & DELETE',
'''public function update($id)
{
    $barang = $this->model->find($id);
    $data   = $this->request->getPost();
    // ... proses ganti / hapus gambar lama ...
    $this->model->update($id, $data);
    return redirect()->to('barang')->with('success', 'Barang berhasil diperbarui.');
}

public function delete($id)
{
    $barang = $this->model->find($id);
    if (!empty($barang['gambar'])) {
        @unlink(FCPATH . 'uploads/barang/' . $barang['gambar']);
    }
    $this->model->delete($id);
    return redirect()->to('barang')->with('success', 'Barang berhasil dihapus.');
}''')

h2('b. Penjelasan')
bullet('Create (store): data dari form disimpan dengan $this->model->save(). Gambar diunggah melalui handleUploadGambar() yang memvalidasi format (JPG/PNG/WEBP/GIF) dan ukuran maksimal 2 MB.')
bullet('Read (index): menampilkan daftar barang dengan join kategori, dilengkapi pencarian (nama/kode), filter kategori, dan paginasi 10 data per halaman.')
bullet('Update (update): mengubah data barang berdasarkan id, sekaligus mengganti atau menghapus gambar lama bila diperlukan.')
bullet('Delete (delete): menghapus data barang beserta file gambarnya dari folder uploads.')
bullet('Validasi pada Model memastikan kode barang unik dan seluruh field wajib terisi; generateKode() membuat kode barang otomatis berurutan.')

h2('c. Screenshot')
screenshot('Gambar 4.5.1  Halaman Daftar Barang (Read)')
screenshot('Gambar 4.5.2  Form Tambah / Edit Barang (Create & Update)')

# ============================================================
# 4.6 TRANSAKSI PENJUALAN
# ============================================================
h1('4.6  Membuat Transaksi Penjualan')

para('Fitur transaksi penjualan digunakan untuk mencatat penjualan barang kepada '
     'pembeli. Saat transaksi disimpan, sistem secara otomatis mengurangi stok barang, '
     'menghitung total belanja, uang bayar, dan kembalian, serta menyimpan snapshot harga '
     'beli agar keuntungan dapat dihitung akurat. Seluruh proses dibungkus dalam database '
     'transaction sehingga data tetap konsisten apabila terjadi kegagalan.')

h2('a. Potongan Kode')

code_block('Route — app/Config/Routes.php',
'''$routes->get('penjualan',                'PenjualanController::index');
$routes->get('penjualan/create',         'PenjualanController::create');
$routes->post('penjualan/store',         'PenjualanController::store');
$routes->get('penjualan/show/(:num)',    'PenjualanController::show/$1');
$routes->get('penjualan/invoice/(:num)', 'PenjualanController::invoice/$1');''')

code_block('Model — nomor transaksi otomatis (PenjualanModel.php)',
'''// Format: PJ-YYYYMMDD-XXX
public function generateNoTransaksi()
{
    $prefix = 'PJ-' . date('Ymd') . '-';
    $last   = $this->like('no_transaksi', $prefix, 'after')
                   ->orderBy('id', 'DESC')->first();
    if (!$last) return $prefix . '001';
    $num = (int) substr($last['no_transaksi'], -3);
    return $prefix . str_pad($num + 1, 3, '0', STR_PAD_LEFT);
}''')

code_block('Controller — menyimpan transaksi (store)',
'''public function store()
{
    $db = \\Config\\Database::connect();
    $db->transStart();                          // mulai transaksi DB

    $detail     = $this->request->getPost('detail') ?? [];
    $totalHarga = array_sum(array_map(fn($d) => $d['jumlah'] * $d['harga_jual'], $detail));
    $bayar      = (float) $this->request->getPost('bayar');

    // Validasi stok mencukupi
    $barangModel = new BarangModel();
    foreach ($detail as $d) {
        $brg = $barangModel->find($d['id_barang']);
        if ($brg && $brg['stok'] < $d['jumlah']) {
            $db->transRollback();
            return redirect()->back()->withInput()
                ->with('error', "Stok {$brg['nama_barang']} tidak cukup.");
        }
    }

    $header = [
        'no_transaksi' => $this->request->getPost('no_transaksi'),
        'tanggal_jual' => $this->request->getPost('tanggal_jual'),
        'nama_pembeli' => $this->request->getPost('nama_pembeli') ?: 'Umum',
        'total_harga'  => $totalHarga,
        'bayar'        => $bayar,
        'kembalian'    => $bayar - $totalHarga,
    ];
    $this->model->insert($header);
    $id = $this->model->getInsertID();

    $this->detailModel->simpanDanKurangiStok($detail, $id);  // simpan detail + kurangi stok
    $db->transComplete();

    return redirect()->to('penjualan/show/' . $id)
        ->with('success', 'Transaksi berhasil disimpan.');
}''')

code_block('Model — simpan detail & kurangi stok otomatis (DetailPenjualanModel.php)',
'''public function simpanDanKurangiStok(array $detail, int $id_penjualan)
{
    $barangModel = new BarangModel();
    $db          = \\Config\\Database::connect();

    foreach ($detail as $item) {
        $barang = $barangModel->find($item['id_barang']);

        $item['id_penjualan'] = $id_penjualan;
        $item['subtotal']     = (int)$item['jumlah'] * (float)$item['harga_jual'];
        $item['harga_beli']   = $barang ? (float) $barang['harga_beli'] : 0; // snapshot
        $this->insert($item);

        // Kurangi stok barang
        $db->table('barang')
           ->where('id', (int) $item['id_barang'])
           ->set('stok', 'stok - ' . (int)$item['jumlah'], false)
           ->update();
    }
}''')

h2('b. Penjelasan')
bullet('Nomor transaksi dibuat otomatis dengan format PJ-YYYYMMDD-XXX melalui generateNoTransaksi().')
bullet('Pada method store(), proses diawali transStart() dan diakhiri transComplete() agar header dan detail tersimpan secara atomik (database transaction).')
bullet('Sebelum menyimpan, sistem memeriksa apakah stok setiap barang mencukupi; jika tidak, transaksi dibatalkan (transRollback) dengan pesan error.')
bullet('Total harga, dan kembalian (bayar - total) dihitung otomatis di sisi server, sementara di sisi tampilan dihitung real-time menggunakan Alpine.js.')
bullet('simpanDanKurangiStok() menyimpan tiap baris detail, menyimpan snapshot harga beli untuk perhitungan keuntungan, lalu mengurangi stok barang secara otomatis.')

h2('c. Screenshot')
screenshot('Gambar 4.6.1  Form Transaksi Penjualan Baru')
screenshot('Gambar 4.6.2  Detail / Invoice Transaksi Penjualan')

# ============================================================
# 4.7 LAPORAN
# ============================================================
h1('4.7  Membuat Laporan')

para('Fitur laporan menyajikan rekapitulasi data dalam tiga jenis: Laporan Penjualan, '
     'Laporan Stok Barang, dan Laporan Barang Masuk. Setiap laporan dapat difilter '
     'berdasarkan periode tanggal/kategori, ditampilkan di layar, serta diekspor ke '
     'format PDF (menggunakan pustaka Dompdf) dan Excel (menggunakan PhpSpreadsheet). '
     'Laporan penjualan juga menghitung total keuntungan dari selisih harga jual dan '
     'harga beli.')

h2('a. Potongan Kode')

code_block('Route — app/Config/Routes.php',
'''$routes->get('laporan',                 'LaporanController::index');
$routes->get('laporan/penjualan',       'LaporanController::penjualan');
$routes->get('laporan/penjualan/pdf',   'LaporanController::penjualanPdf');
$routes->get('laporan/penjualan/excel', 'LaporanController::penjualanExcel');
$routes->get('laporan/stok',            'LaporanController::stok');
$routes->get('laporan/barang-masuk',    'LaporanController::barangMasuk');''')

code_block('Controller — menampilkan laporan penjualan (LaporanController.php)',
'''public function penjualan()
{
    $dari   = $this->request->getGet('dari')   ?? date('Y-m-01');
    $sampai = $this->request->getGet('sampai') ?? date('Y-m-d');

    $data = $this->_queryPenjualan($dari, $sampai);

    return view('laporan/penjualan', [
        'rows'             => $data,
        'dari'             => $dari,
        'sampai'           => $sampai,
        'total'            => array_sum(array_column($data, 'total_harga')),
        'total_keuntungan' => array_sum(array_column($data, 'total_keuntungan')),
    ]);
}''')

code_block('Controller — query data + keuntungan',
'''private function _queryPenjualan(string $dari, string $sampai): array
{
    return \\Config\\Database::connect()->query("
        SELECT p.*,
            COALESCE(SUM((dp.harga_jual - dp.harga_beli) * dp.jumlah), 0) AS total_keuntungan
        FROM penjualan p
        LEFT JOIN detail_penjualan dp ON dp.id_penjualan = p.id
        WHERE p.tanggal_jual >= ? AND p.tanggal_jual <= ?
        GROUP BY p.id
        ORDER BY p.tanggal_jual ASC
    ", [$dari, $sampai])->getResultArray();
}''')

code_block('Controller — ekspor PDF dengan Dompdf',
'''public function penjualanPdf()
{
    $dari   = $this->request->getGet('dari')   ?? date('Y-m-01');
    $sampai = $this->request->getGet('sampai') ?? date('Y-m-d');
    $rows   = $this->_queryPenjualan($dari, $sampai);

    $html = view('laporan/pdf/penjualan', compact('rows', 'dari', 'sampai'));
    $this->_streamPdf($html, 'Laporan_Penjualan_' . $dari . '_' . $sampai, 'landscape');
}

private function _streamPdf(string $html, string $filename, string $orientation = 'portrait')
{
    $options = new Options();
    $options->set('isHtml5ParserEnabled', true);
    $dompdf = new Dompdf($options);
    $dompdf->loadHtml($html);
    $dompdf->setPaper('A4', $orientation);
    $dompdf->render();
    $dompdf->stream($filename . '.pdf', ['Attachment' => true]);
}''')

code_block('Controller — ekspor Excel dengan PhpSpreadsheet (potongan)',
'''$ss    = new Spreadsheet();
$sheet = $ss->getActiveSheet()->setTitle('Penjualan');

$headers = ['No', 'No Transaksi', 'Tanggal', 'Nama Pembeli',
            'Total Harga', 'Bayar', 'Kembalian', 'Keuntungan'];
foreach ($headers as $i => $h) {
    $sheet->setCellValue($cols[$i] . '4', $h);
}

$row = 5;
foreach ($rows as $no => $r) {
    $sheet->setCellValue('A' . $row, $no + 1);
    $sheet->setCellValue('B' . $row, $r['no_transaksi']);
    $sheet->setCellValue('E' . $row, (float) $r['total_harga']);
    $row++;
}

$writer = new Xlsx($ss);
$writer->save('php://output');   // unduh file .xlsx''')

h2('b. Penjelasan')
bullet('Halaman index laporan menampilkan tiga pilihan kartu: Laporan Penjualan, Laporan Stok Barang, dan Laporan Barang Masuk.')
bullet('Setiap laporan menerima parameter filter (dari, sampai, atau id_kategori); jika kosong, default-nya adalah awal bulan s/d hari ini.')
bullet('Data laporan penjualan diambil melalui query SQL yang sekaligus menghitung total keuntungan dari (harga_jual - harga_beli) x jumlah.')
bullet('Ekspor PDF menggunakan Dompdf: view HTML laporan di-render menjadi berkas PDF A4 dan diunduh otomatis.')
bullet('Ekspor Excel menggunakan PhpSpreadsheet: data ditulis ke sel-sel spreadsheet lengkap dengan judul, header berwarna, format angka, dan baris total, lalu diunduh sebagai berkas .xlsx.')

h2('c. Screenshot')
screenshot('Gambar 4.7.1  Halaman Pilihan Laporan')
screenshot('Gambar 4.7.2  Tampilan Laporan Penjualan (filter periode)')
screenshot('Gambar 4.7.3  Hasil Ekspor PDF / Excel')

# ---------- Simpan ----------
out = r'c:\laragon\www\Manajemen-Inventori\docs\BAB IV - 4.4 sd 4.7 Manajemen Inventori.docx'
doc.save(out)
print('SAVED:', out)
